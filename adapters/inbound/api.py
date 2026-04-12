import io
import os
import re
import uuid
import shutil
from datetime import date

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles

from adapters.outbound.pdf_adapter import PyMuPDFAdapter
from adapters.outbound.embedding_adapter import FAISSEmbeddingAdapter
from adapters.outbound.llm_adapter import QwenAdapter
from adapters.outbound.history_adapter import FileConversationHistory
from application.ask_question import AskQuestionUseCase

UPLOADS_DIR = "uploads"
os.makedirs(UPLOADS_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Adapter singletons — heavy models are loaded once at startup
# ---------------------------------------------------------------------------
_pdf = PyMuPDFAdapter()
_emb = FAISSEmbeddingAdapter()
_llm = QwenAdapter()
_history = FileConversationHistory()
_use_case = AskQuestionUseCase(_pdf, _emb, _llm, _history)

# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------
app = FastAPI(title="PdfTool")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# --- Session management ----------------------------------------------------

@app.post("/api/sessions")
def create_session():
    return {"session_id": str(uuid.uuid4())}


@app.get("/api/sessions")
def list_sessions():
    convs = _history.list_sessions()
    return {
        "sessions": [
            {"id": c.id, "label": c.label or c.id[:8], "message_count": len(c.messages)}
            for c in convs
        ]
    }


# --- PDF upload ------------------------------------------------------------

@app.post("/api/sessions/{session_id}/upload")
async def upload_pdf(session_id: str, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted")
    path = os.path.join(UPLOADS_DIR, f"{session_id}.pdf")
    with open(path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    _history.set_label(session_id, file.filename)

    # Build and persist the FAISS index immediately after upload so that
    # subsequent /ask calls load from disk instead of re-encoding every time.
    pages = _pdf.extract_pages(path)
    _emb.index_session(session_id, pages)

    return {"status": "ok", "filename": file.filename}


# --- Q&A -------------------------------------------------------------------

@app.post("/api/sessions/{session_id}/ask")
async def ask(session_id: str, query: str = Form(...)):
    pdf_path = os.path.join(UPLOADS_DIR, f"{session_id}.pdf")
    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="No PDF uploaded for this session")
    result = _use_case.execute(session_id, query, pdf_path)
    return {
        "answer": result.answer,
        "keywords": result.keywords,
        "source_pages": result.source_pages,
    }


# --- Export (PDF / DOCX) ---------------------------------------------------

@app.post("/api/sessions/{session_id}/export")
async def export_document(
    session_id: str,
    title: str = Form(...),
    text: str = Form(...),
    source_pages: str = Form(""),
    format: str = Form("pdf"),
):
    pages = [int(p) for p in source_pages.split(",") if p.strip().isdigit()]
    fmt = format.lower()

    if fmt == "docx":
        buf = _make_docx(title, text, pages)
        filename = "answer.docx"
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        buf = _make_pdf(title, text, pages)
        filename = "answer.pdf"
        media_type = "application/pdf"

    return StreamingResponse(
        buf,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# --- History ---------------------------------------------------------------

@app.get("/api/sessions/{session_id}/history")
def get_history(session_id: str):
    conv = _history.get_or_create(session_id)
    return {
        "id": conv.id,
        "label": conv.label,
        "messages": [{"role": m.role, "content": m.content} for m in conv.messages],
    }


# --- Serve frontend — must be registered LAST so API routes take priority ---
app.mount("/", StaticFiles(directory="frontend/dist", html=True), name="frontend")


# ---------------------------------------------------------------------------
# Document generation helpers
# ---------------------------------------------------------------------------

def _rl_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _rl_inline(text: str) -> str:
    """Convert **bold** *italic* `code` to ReportLab XML tags."""
    text = _rl_escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`(.+?)`", r'<font name="Courier">\1</font>', text)
    return text


def _make_pdf(title: str, text: str, source_pages: list) -> io.BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "Q", parent=styles["Normal"],
        fontSize=16, fontName="Helvetica-Bold",
        spaceAfter=6, textColor=colors.HexColor("#1e1e2e"),
    )
    body_style = ParagraphStyle(
        "B", parent=styles["Normal"],
        fontSize=11, leading=17, spaceAfter=4,
        textColor=colors.HexColor("#2d2d3f"),
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Normal"],
        fontSize=13, fontName="Helvetica-Bold",
        spaceBefore=10, spaceAfter=4,
        textColor=colors.HexColor("#1e1e2e"),
    )
    h3_style = ParagraphStyle(
        "H3", parent=styles["Normal"],
        fontSize=12, fontName="Helvetica-BoldOblique",
        spaceBefore=8, spaceAfter=3,
        textColor=colors.HexColor("#2d2d3f"),
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body_style,
        leftIndent=18, firstLineIndent=-10,
    )
    code_style = ParagraphStyle(
        "Code", parent=styles["Normal"],
        fontSize=9.5, fontName="Courier",
        backColor=colors.HexColor("#f4f4f8"),
        leftIndent=12, rightIndent=12,
        spaceAfter=4, spaceBefore=4, leading=14,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#888899"), leading=14,
    )

    story = [
        Paragraph(_rl_escape(title), title_style),
        Spacer(1, 0.15 * inch),
        HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#ccccdd")),
        Spacer(1, 0.15 * inch),
    ]

    in_code = False
    code_buf = []
    for line in text.split("\n"):
        if line.startswith("```"):
            if in_code:
                story.append(Paragraph(_rl_escape("\n".join(code_buf)), code_style))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.06 * inch))
            continue

        if stripped.startswith("## "):
            story.append(Paragraph(_rl_escape(stripped[3:]), h2_style))
        elif stripped.startswith("### "):
            story.append(Paragraph(_rl_escape(stripped[4:]), h3_style))
        elif stripped.startswith("# "):
            story.append(Paragraph(_rl_escape(stripped[2:]), h2_style))
        elif re.match(r"^[-*] ", stripped):
            story.append(Paragraph("•  " + _rl_inline(stripped[2:]), bullet_style))
        elif re.match(r"^\d+\. ", stripped):
            content = re.sub(r"^\d+\. ", "", stripped)
            story.append(Paragraph("•  " + _rl_inline(content), bullet_style))
        else:
            story.append(Paragraph(_rl_inline(stripped), body_style))

    if in_code and code_buf:
        story.append(Paragraph(_rl_escape("\n".join(code_buf)), code_style))

    pages_str = ", ".join(f"p.{p}" for p in source_pages) if source_pages else "—"
    story += [
        Spacer(1, 0.25 * inch),
        HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#ccccdd")),
        Spacer(1, 0.08 * inch),
        Paragraph(f"Source pages: {pages_str}  ·  Generated {date.today()}", meta_style),
    ]

    doc.build(story)
    buf.seek(0)
    return buf


def _make_docx(title: str, text: str, source_pages: list) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph()

    in_code = False
    code_lines = []
    for line in text.split("\n"):
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^[-*] ", stripped):
            _docx_inline(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            _docx_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\. ", "", stripped))
        else:
            _docx_inline(doc.add_paragraph(), stripped)

    if in_code and code_lines:
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)

    doc.add_paragraph()
    pages_str = ", ".join(f"p.{p}" for p in source_pages) if source_pages else "—"
    meta = doc.add_paragraph(f"Source pages: {pages_str}  ·  Generated {date.today()}")
    if meta.runs:
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_inline(para, text: str) -> None:
    """Add runs with **bold** *italic* `code` handling to a paragraph."""
    from docx.shared import Pt

    parts = re.split(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            para.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        else:
            para.add_run(part)
